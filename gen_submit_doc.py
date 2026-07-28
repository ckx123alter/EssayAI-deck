"""
生成提交文档 — 告诉面试官如何访问在线演示
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(3); s.bottom_margin = Cm(3)
    s.left_margin = Cm(3); s.right_margin = Cm(3)

style = doc.styles['Normal']
style.font.name = '微软雅黑'; style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.8

def p(text, bold=False, size=None, color=None, align=None):
    par = doc.add_paragraph()
    r = par.add_run(text); r.font.name = '微软雅黑'
    if bold: r.bold = True
    if size: r.font.size = Pt(size)
    if color: r.font.color.rgb = RGBColor(*color)
    if align is not None: par.alignment = align

# 标题
p("EssayAI 产品演示 · 访问说明", bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER)
p("AI 产品经理校招面试 — 产品开放麦", size=12, color=(120,120,120), align=WD_ALIGN_PARAGRAPH.CENTER)
p("")

# 线上演示链接
p("\U0001F517 在线演示地址", bold=True, size=14)
p("https://ckx123alter.github.io/EssayAI-deck/", size=13, color=(74,85,104))
p("")
p("请复制上方链接到浏览器地址栏打开（推荐 Chrome / Edge）。"
  "建议按 F11 全屏浏览，效果等同于 PPT 幻灯片播放。", size=10, color=(120,120,120))
p("")

# 操作说明
p("\U0001F4A1 如何操作", bold=True, size=14)
tips = [
    "点击任意页面缩略图 → 进入演示模式",
    "键盘 ← → 或点击屏幕左右两侧 → 翻页",
    "按 Esc 键 → 返回概览页，可重新选择页面",
    "第 5 页「产品效果」包含演示视频，点击 ▶ 按钮即可播放",
]
for t in tips:
    p(f"  ·  {t}", size=11)
p("")

# 备选方案
p("\U0001F4E6 备选方案", bold=True, size=14)
p("如网络原因无法访问线上链接，可联系我通过其他方式发送演示文件。", size=11)
p("")

# 联系方式
p("\U0001F4DE 联系方式", bold=True, size=14)
p("陈可莘  ·  19821742260  ·  530946401@qq.com", size=11)
p("")
p("产品：EssayAI — 让每个人拥有顶级文书", size=10, color=(150,150,150))

out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "EssayAI_演示访问说明_提交用.docx")
doc.save(out)
print(f"Done: {out}")
